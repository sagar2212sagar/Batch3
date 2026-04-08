from pyspark.sql.functions import udf
from pyspark.sql.types import BooleanType, StringType
import re
import io

@udf(returnType=BooleanType())
def is_valid_email(email):
    """
    This function checks if the given email address has a valid format using regex.
    Returns True if valid, False otherwise.
    """
    pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
    if email is None:
        return False
    return re.match(pattern, email) is not None

@udf(returnType=StringType())
def parse_pdf_content(content):
    """
    Extracts text content from PDF binary data.
    """
    from PyPDF2 import PdfReader
    
    if content is None:
        return None
    
    try:
        # Wrap binary content in BytesIO for PyPDF2
        pdf_file = io.BytesIO(content)
        reader = PdfReader(pdf_file)
        
        text = ""
        for page in reader.pages:
            text += page.extract_text()
        
        return text
    except Exception as e:
        return f"Error parsing PDF: {str(e)}"
    


@udf(returnType=StringType())
def parse_content(path,content):
    """
    Extracts text content from DOCX binary data.
    """
    from docx import Document
    
    if content is None:
        return None
    
    try:
        # Wrap binary content in BytesIO for python-docx
        docx_file = io.BytesIO(content)
        doc = Document(docx_file)
        text = ""
        for para in doc.paragraphs:
            text += para.text
        return text
    except Exception as e:
        return f"Error parsing DOCX: {str(e)}"
    

@udf(returnType=StringType())
def parse_pdf_content(content):
    """
    Extracts text content from PDF binary data.
    """
    from PyPDF2 import PdfReader
    
    if content is None:
        return None
    
    try:
        # Wrap binary content in BytesIO for PyPDF2
        pdf_file = io.BytesIO(content)
        reader = PdfReader(pdf_file)
        
        text = ""
        for page in reader.pages:
            text += page.extract_text()
        
        return text
    except Exception as e:
        return f"Error parsing PDF: {str(e)}"
    
@udf(returnType=StringType())
def parse_docx_content(content):
    """
    Extracts text content from DOCX binary data.
    """
    from docx import Document
    
    if content is None:
        return None
    
    try:
        # Wrap binary content in BytesIO for python-docx
        docx_file = io.BytesIO(content)
        doc = Document(docx_file)
        text = ""
        for para in doc.paragraphs:
            text += para.text
        return text
    except Exception as e:
        return f"Error parsing DOCX: {str(e)}"
    